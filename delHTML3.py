import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# 创建一个目录来存储下载的图片
image_dir = 'images'
if not os.path.exists(image_dir):
    os.makedirs(image_dir)

# 打开文件进行读取
with open('que.txt', 'r', encoding='utf-8') as file:
    file_contents = file.read()

# 创建Beautiful Soup对象来解析文本
soup = BeautifulSoup(file_contents, 'html.parser')

# 创建一个新的Word文档
doc = Document()

# 使用类名选择器查找所有问题块
question_blocks = soup.find_all("div", class_='questionItem___q6Hgu')

for question_block in question_blocks:
    # 提取问题文本
    question_text = question_block.text.strip()
    doc.add_heading("问题:", level=1)
    doc.add_paragraph(question_text)

    # 提取答案选项文本
    answer_elements = question_block.select('label.ant-radio-wrapper')
    for answer_element in answer_elements:
        option = answer_element.select('span.font16')[0].text.strip()
        text = answer_element.select('div.renderHtml___UerV1')[0].text.strip()
        doc.add_heading("选项: {}".format(option), level=2)
        doc.add_paragraph("答案: {}".format(text))
    
    # 提取图片链接
    image_elements = question_block.select('img')
    for i, image_element in enumerate(image_elements):
        image_url = image_element['src']
        # 下载图片并保存到本地文件夹
        response = requests.get(image_url)
        if response.status_code == 200:
            image_filename = os.path.join(image_dir, 'image_{}.png'.format(i))
            with open(image_filename, 'wb') as image_file:
                image_file.write(response.content)
            # 将图片插入到Word文档中
            doc.add_heading("图片链接:", level=2)
            doc.add_paragraph(image_url)
            doc.add_heading("图片:", level=2)
            doc.add_picture(image_filename, width=Inches(4.0))
    
    # 添加分页符来分隔不同的问题
    doc.add_page_break()
    doc.save('output.docx')
    for i in range(1,200):
        print(i)

# 保存Word文档
doc.save('output.docx')
print("已保存")

